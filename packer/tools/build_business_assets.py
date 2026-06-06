from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.formatting.rule import FormulaRule
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.worksheet.datavalidation import DataValidation
from openpyxl.worksheet.table import Table, TableStyleInfo
from openpyxl.utils import get_column_letter


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_PATH = ROOT / "templates" / "Cstar货物导入模板.xlsx"
MANUAL_PATH = ROOT / "docs" / "喜事达Cstar装箱软件用户手册.docx"

ACCENT = "EB2325"
NAVY = "111111"
LIGHT = "FFF1F1"
MID = "F6DADA"
WARN = "FFF2CC"
ERROR = "FCE4D6"
GRAY = "F4F6F8"
BORDER = "D9E0E7"


def style_range(ws, cell_range, fill=None, font=None, align=None, border=True):
    side = Side(style="thin", color=BORDER)
    for row in ws[cell_range]:
        for cell in row:
            if fill:
                cell.fill = fill
            if font:
                cell.font = font
            if align:
                cell.alignment = align
            if border:
                cell.border = Border(left=side, right=side, top=side, bottom=side)


def add_note(ws, row, text):
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=15)
    cell = ws.cell(row, 1, text)
    cell.fill = PatternFill("solid", fgColor=LIGHT)
    cell.font = Font(color=NAVY, italic=True)
    cell.alignment = Alignment(wrap_text=True, vertical="center")


def build_excel_template():
    TEMPLATE_PATH.parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook()
    ws = wb.active
    ws.title = "货物录入"
    ws.freeze_panes = "A9"
    ws.sheet_view.showGridLines = False

    meta = wb.create_sheet("填写说明")
    refs = wb.create_sheet("字段字典")
    containers = wb.create_sheet("箱型参考")
    examples = wb.create_sheet("示例数据")

    title_fill = PatternFill("solid", fgColor=NAVY)
    header_fill = PatternFill("solid", fgColor=ACCENT)
    sub_fill = PatternFill("solid", fgColor=MID)
    white_bold = Font(color="FFFFFF", bold=True)
    title_font = Font(color="FFFFFF", bold=True, size=16)
    header_font = Font(color="FFFFFF", bold=True)
    label_font = Font(color=NAVY, bold=True)
    normal_align = Alignment(vertical="center", wrap_text=True)
    center_align = Alignment(horizontal="center", vertical="center", wrap_text=True)

    ws.merge_cells("A1:O1")
    ws["A1"] = "喜事达（Cstar）装箱软件 - 货物导入模板"
    ws["A1"].fill = title_fill
    ws["A1"].font = title_font
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 30

    project_rows = [
        ("任务名称", "请填写"),
        ("客户/订单", "请填写"),
        ("操作员", "请填写"),
        ("装箱日期", "YYYY-MM-DD"),
    ]
    for i, (label, value) in enumerate(project_rows, start=3):
        ws.cell(i, 1, label)
        ws.cell(i, 2, value)
    style_range(ws, "A3:B6", fill=PatternFill("solid", fgColor=GRAY), font=label_font, align=normal_align)
    style_range(ws, "B3:B6", fill=PatternFill("solid", fgColor="FFFFFF"), font=Font(color=NAVY), align=normal_align)

    add_note(ws, 7, "只需填写第 9 行以下的货物数据。尺寸单位为毫米，重量单位为千克。请保留第 8 行字段名。")

    headers = ["名称", "长", "宽", "高", "重量", "数量", "分组", "可旋转", "可倾斜", "可堆叠", "优先级", "备注", "校验", "体积", "箱门校验"]
    ws.append(headers)
    ws.row_dimensions[8].height = 24
    style_range(ws, "A8:O8", fill=header_fill, font=header_font, align=center_align)

    sample_rows = [
        ["家电纸箱 A", 1200, 800, 900, 95, 18, "上海", "是", "是", "是", 3, "可旋转可堆叠", "", "", ""],
        ["托盘货 B", 1100, 1000, 1200, 380, 10, "上海", "是", "否", "否", 3, "重货，不允许倾斜/堆叠", "", "", ""],
        ["长条设备 C", 2400, 600, 700, 210, 8, "宁波", "是", "是", "是", 2, "优先级中", "", "", ""],
        ["配件箱 D", 600, 500, 450, 35, 40, "宁波", "是", "是", "是", 1, "小包装", "", "", ""],
    ]
    for row in sample_rows:
        ws.append(row)

    for row in range(9, 209):
        ws.cell(row, 13, f'=IF(OR(A{row}="",B{row}="",C{row}="",D{row}="",E{row}="",F{row}="",G{row}=""),"待填写",IF(AND(B{row}>0,C{row}>0,D{row}>0,E{row}>=0,F{row}>=1),"通过","需检查"))')
        ws.cell(row, 14, f'=IF(AND(B{row}>0,C{row}>0,D{row}>0),B{row}*C{row}*D{row}/1000000000,"")')
        ws.cell(row, 15, f'=IF(OR(C{row}="",D{row}=""),"待填写",IF(OR(AND(C{row}<=2340,D{row}<=2585),AND(D{row}<=2340,C{row}<=2585)),"通过","检查箱门"))')

    widths = [24, 11, 11, 11, 11, 10, 20, 12, 12, 12, 10, 28, 12, 12, 14]
    for idx, width in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(idx)].width = width
    for row in ws.iter_rows(min_row=9, max_row=208, max_col=15):
        for cell in row:
            cell.alignment = center_align if cell.column in [2, 3, 4, 5, 6, 8, 9, 10, 11, 13, 14, 15] else normal_align
            cell.border = Border(left=Side(style="thin", color=BORDER), right=Side(style="thin", color=BORDER), top=Side(style="thin", color=BORDER), bottom=Side(style="thin", color=BORDER))
    ws.column_dimensions["L"].width = 28
    ws.column_dimensions["A"].width = 24

    table = Table(displayName="CargoInputTable", ref="A8:O208")
    table.tableStyleInfo = TableStyleInfo(name="TableStyleMedium4", showFirstColumn=False, showLastColumn=False, showRowStripes=True, showColumnStripes=False)
    ws.add_table(table)

    for col in ["B", "C", "D"]:
        dv = DataValidation(type="whole", operator="greaterThan", formula1="0", allow_blank=False)
        dv.error = "尺寸必须为大于 0 的整数，单位毫米。"
        dv.errorTitle = "尺寸填写错误"
        ws.add_data_validation(dv)
        dv.add(f"{col}9:{col}208")
    dv_weight = DataValidation(type="decimal", operator="greaterThanOrEqual", formula1="0", allow_blank=False)
    dv_weight.error = "重量不能为负数，单位千克。"
    ws.add_data_validation(dv_weight)
    dv_weight.add("E9:E208")
    dv_qty = DataValidation(type="whole", operator="greaterThanOrEqual", formula1="1", allow_blank=False)
    dv_qty.error = "数量必须为正整数。"
    ws.add_data_validation(dv_qty)
    dv_qty.add("F9:F208")
    dv_bool = DataValidation(type="list", formula1='"是,否"', allow_blank=False)
    ws.add_data_validation(dv_bool)
    dv_bool.add("H9:J208")
    dv_priority = DataValidation(type="list", formula1='"3,2,1"', allow_blank=False)
    ws.add_data_validation(dv_priority)
    dv_priority.add("K9:K208")

    ws.conditional_formatting.add("M9:M208", FormulaRule(formula=['$M9="需检查"'], fill=PatternFill("solid", fgColor=ERROR)))
    ws.conditional_formatting.add("M9:M208", FormulaRule(formula=['$M9="待填写"'], fill=PatternFill("solid", fgColor=WARN)))
    ws.conditional_formatting.add("M9:M208", FormulaRule(formula=['$M9="通过"'], fill=PatternFill("solid", fgColor=LIGHT)))
    ws.conditional_formatting.add("O9:O208", FormulaRule(formula=['$O9="检查箱门"'], fill=PatternFill("solid", fgColor=ERROR)))
    ws.conditional_formatting.add("A9:K208", FormulaRule(formula=['$M9="需检查"'], fill=PatternFill("solid", fgColor=ERROR)))

    # Instructions sheet
    meta.sheet_view.showGridLines = False
    meta.merge_cells("A1:F1")
    meta["A1"] = "填写说明"
    meta["A1"].fill = title_fill
    meta["A1"].font = title_font
    meta["A1"].alignment = Alignment(horizontal="center")
    meta_rows = [
        ["步骤", "说明"],
        ["1", "在“货物录入”表第 9 行开始填写货物；保留第 8 行字段名。"],
        ["2", "尺寸单位为毫米，重量单位为千克；数量必须为正整数。"],
        ["3", "可旋转、可倾斜、可堆叠使用下拉选择“是”或“否”。"],
        ["4", "优先级：3=高，2=中，1=低；高优先级会优先装载。"],
        ["5", "保存后在喜事达（Cstar）装箱软件中点击“导入货物”，选择本工作簿。"],
        ["6", "如果校验列显示“需检查”，请先修正该行再导入。"],
    ]
    for row in meta_rows:
        meta.append(row)
    style_range(meta, "A3:B8", align=normal_align)
    style_range(meta, "A2:B2", fill=header_fill, font=header_font, align=center_align)
    meta.column_dimensions["A"].width = 12
    meta.column_dimensions["B"].width = 78

    # Field dictionary
    refs.sheet_view.showGridLines = False
    refs.append(["字段", "含义", "是否必填", "填写规则", "示例"])
    field_rows = [
        ["名称", "货物名称", "是", "文本，不建议重复过多简称", "家电纸箱 A"],
        ["长", "长度", "是", "大于 0，单位毫米", "1200"],
        ["宽", "宽度", "是", "大于 0，单位毫米", "800"],
        ["高", "高度", "是", "大于 0，单位毫米", "900"],
        ["重量", "单件重量", "是", "大于等于 0，单位千克", "95"],
        ["数量", "数量", "是", "正整数", "18"],
        ["分组", "分组/目的地", "是", "客户、订单、目的港或卸货点", "上海"],
        ["可旋转", "是否可旋转", "是", "是或否", "是"],
        ["可倾斜", "是否可倾斜", "是", "是表示可侧放；否表示保持原高度", "否"],
        ["可堆叠", "是否可堆叠", "是", "是或否", "否"],
        ["优先级", "优先级", "是", "3=高，2=中，1=低", "3"],
    ]
    for row in field_rows:
        refs.append(row)
    style_range(refs, "A1:E1", fill=header_fill, font=header_font, align=center_align)
    style_range(refs, "A2:E12", align=normal_align)
    for col, width in zip("ABCDE", [16, 16, 12, 42, 18]):
        refs.column_dimensions[col].width = width

    # Container reference
    containers.append(["箱型", "内长毫米", "内宽毫米", "内高毫米", "箱门宽毫米", "箱门高毫米", "最大载重千克"])
    for row in [
        ["20GP", 5898, 2352, 2393, 2340, 2280, 28200],
        ["40GP", 12032, 2352, 2393, 2340, 2280, 28600],
        ["40HQ", 12032, 2352, 2698, 2340, 2585, 28600],
        ["45HQ", 13556, 2352, 2698, 2340, 2585, 29500],
    ]:
        containers.append(row)
    style_range(containers, "A1:G1", fill=header_fill, font=header_font, align=center_align)
    style_range(containers, "A2:G5", align=center_align)
    for col in "ABCDEFG":
        containers.column_dimensions[col].width = 15

    # Example data
    examples.append(headers[:11])
    for row in sample_rows:
        examples.append(row[:11])
    style_range(examples, "A1:K1", fill=header_fill, font=header_font, align=center_align)
    style_range(examples, "A2:K5", align=center_align)
    for col in "ABCDEFGHIJK":
        examples.column_dimensions[col].width = 15
    examples.column_dimensions["A"].width = 24

    for sheet in [ws, meta, refs, containers, examples]:
        for row in sheet.iter_rows():
            for cell in row:
                cell.alignment = cell.alignment.copy(vertical="center", wrap_text=True)

    wb.save(TEMPLATE_PATH)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")
    tc_pr.append(tc_w)


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(0x98, 0x23, 0x25)
    return paragraph


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    set_cell_width(cell, 9360)
    p = cell.paragraphs[0]
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor(0x13, 0x20, 0x2C)
    p.add_run(f"\n{text}")


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, "E8EEF5")
        set_cell_width(cell, widths[idx])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            set_cell_width(cells[idx], widths[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return table


def build_manual():
    MANUAL_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.25
    for style_name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 12)]:
        styles[style_name].font.name = "Calibri"
        styles[style_name].font.size = Pt(size)
        styles[style_name].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        styles[style_name].paragraph_format.space_before = Pt(12)
        styles[style_name].paragraph_format.space_after = Pt(6)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("喜事达（Cstar）装箱软件用户手册")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
    subtitle = doc.add_paragraph("Cstar 装箱软件操作指南")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.color.rgb = RGBColor(0x60, 0x70, 0x80)

    add_callout(
        doc,
        "适用范围",
        "本手册面向业务录单、仓库装柜和管理复核人员。软件用于快速生成装箱参考方案，正式发运前仍需现场复核承重、绑扎、法规、箱况和实际包装。",
    )

    add_heading(doc, "1. 快速开始", 1)
    for step in [
        "解压软件包，打开首页文件。",
        "填写任务名称、客户/订单、操作员和装箱日期。",
        "点击“导入货物”，选择逗号表、工作簿或文字文档。",
        "选择箱型、最大箱数和约束规则，点击“自动装箱”。",
        "按箱号、视图模式和分层高度检查装载结果。",
        "在“装箱报告”页复制报告、导出表格或打印存档。",
    ]:
        add_number(doc, step)

    add_heading(doc, "2. 专业模板填写规范", 1)
    add_table(
        doc,
        ["字段", "说明", "规则"],
        [
            ["名称", "货物名称", "必填文本，例如“家电纸箱 A”"],
            ["长、宽、高", "尺寸", "单位毫米，必须大于 0"],
            ["重量", "单件重量", "单位千克，不能为负数"],
            ["数量", "数量", "必须为正整数"],
            ["分组", "分组/目的地", "客户、订单、目的港或卸货点"],
            ["可旋转", "是否可旋转", "是或否"],
            ["可倾斜", "是否可倾斜", "是表示可侧放，否表示保持原高度"],
            ["可堆叠", "是否可堆叠", "是或否"],
            ["优先级", "装载优先级", "3=高，2=中，1=低"],
        ],
        [1600, 2500, 5260],
    )
    add_bullet(doc, "不要修改模板第 8 行字段名，软件依靠这些字段识别货物。")
    add_bullet(doc, "如果校验列显示“需检查”，请先修正后再导入。")
    add_bullet(doc, "文字文档也可以导入，但建议日常业务优先使用专业模板。")

    add_heading(doc, "3. 自动装箱、多箱分装与装载顺序", 1)
    add_bullet(doc, "最大箱数用于控制可使用的集装箱数量。软件会先装满前一个箱，再把剩余货物分配到后续箱。")
    add_bullet(doc, "约束规则包括重货优先、不可堆叠、左右重量均衡估算。")
    add_bullet(doc, "结果队列会提示未装货物、总体积、总重量、重心和智能建议。")
    add_bullet(doc, "箱门宽高会参与校验，避免货物虽然能摆进箱体但无法通过箱门。")
    add_bullet(doc, "报告会生成建议装载顺序，仓库可按箱号、优先级、重量和箱内位置复核装载顺序。")

    add_heading(doc, "4. 市场对标与新增能力", 1)
    add_table(
        doc,
        ["对标能力", "成熟软件常见做法", "Cstar 当前实现"],
        [
            ["数据导入", "批量导入货物数据", "支持逗号表、工作簿、文字文档，并提供专业模板"],
            ["手动编辑", "拖拽或手动调整货物", "支持选择单件货物并 100 毫米步进微调"],
            ["优先级/分组", "按目的地或优先级管理装载", "新增 group 分组/目的地字段和装载顺序"],
            ["旋转/倾斜限制", "控制货物可否旋转、倾斜、堆叠", "支持可旋转、可倾斜、可堆叠"],
            ["报告与分享", "导出报告并分享方案", "支持表格导出、打印存档、项目文件交接"],
        ],
        [1900, 3600, 3860],
    )

    add_heading(doc, "5. 装载视图、分层与移动", 1)
    add_table(
        doc,
        ["功能", "用途", "操作"],
        [
            ["箱号选择", "查看多箱方案中的某一个箱", "在装载视图右上角选择 1号箱、2号箱等"],
            ["等距/俯视/侧视", "从不同角度检查装载", "点击视图切换按钮"],
            ["分层显示", "小包装较多时减少遮挡", "拖动“分层显示高度”滑杆"],
            ["选择货物", "定位具体单件货物", "在货物下拉框选择货物编号"],
            ["微调移动", "对个别货物做 100 毫米步进调整", "使用前/后/左/右/上/下按钮"],
        ],
        [1600, 3300, 4460],
    )
    add_callout(doc, "移动限制", "手动移动会自动阻止越界和碰撞。若按钮提示“位置不可用”，说明目标位置与其他货物冲突、超出箱体或不符合箱门约束。")

    add_heading(doc, "6. 报告与交接", 1)
    add_bullet(doc, "“保存项目”会保存到当前浏览器。")
    add_bullet(doc, "“导出项目”会生成项目文件，适合同事之间继续编辑同一个任务。")
    add_bullet(doc, "“下载表格”导出每件货物的箱号、坐标、尺寸、重量和旋转方向。")
    add_bullet(doc, "“打印存档”适合归档或交给仓库执行。")

    add_heading(doc, "7. 常见问题", 1)
    add_table(
        doc,
        ["问题", "处理建议"],
        [
            ["导入后没有货物", "确认第一行是字段名，至少包含名称、长、宽、高、重量、数量。"],
            ["货物无法装入", "检查尺寸、重量、箱门宽高、不可堆叠和最大箱数。"],
            ["小货物看不清", "使用分层显示，并在货物下拉框中选择具体货物。"],
            ["同事打开是旧版", "请重新解压最新版压缩包，或刷新浏览器缓存。"],
            ["正式发运是否可直接按软件执行", "不建议直接执行。软件是作业参考，现场仍需复核承重、绑扎和安全要求。"],
        ],
        [2600, 6760],
    )

    footer = section.footer.paragraphs[0]
    footer.text = "喜事达（Cstar）装箱软件用户手册 | 版本 1.2"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(0x60, 0x70, 0x80)

    doc.save(MANUAL_PATH)


if __name__ == "__main__":
    build_excel_template()
    build_manual()
    print(TEMPLATE_PATH)
    print(MANUAL_PATH)
